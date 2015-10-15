# -*- coding: utf-8 -*-
"""
Created on Fri Sep 11 19:36:35 2015

@author: abhishek
"""
import re

#import xl
def convertDoc2Docx(pathName,docName):
    import win32com.client
    wrd= win32com.client.Dispatch("Word.Application")
    #wrd.visible = 0
    a = wrd.Documents.Open(join(pathName,docName))
    docName = docName.split('.')[0]+'.docx'
    a.SaveAs2(join(pathName,docName), FileFormat = 12)
    a.Close()
    return docName
#a.Quit()

hobbiesDict = ['Hobbies', 'Hobby', 'interest']
AchievementDict = ['Achievement', 'award', 'scholar']
CollegeGradDict =	['Indian Institute of Technology','IIT Delhi','IIT Kanpur','IIT Bombay','IIT Kharagpur','IIT Madras','IIT Roorkee','BITS Pilani','IIT BHU','IIT Guwahati','Delhi Techonolgical University','ISM Dhanbad','National Institute of Technology','NIT Tiruchirapalli','VNIT','NIT Surathkal','College of Engineering, Amma University','Birla Institute of Technology, Goa','BIT, Mesra','Manipal Institute of Technology','MIT Pune','College of Engineering, Pune'	,'Vellore Institute of Technology','Veer Jeejamata Technological Institute','Thapar Institute of Technology' , 'institute' , 'engineering', 'college', 'university']
CollegePostGradDict = ['Post Graduation','Indian Institute of Management','IIM Ahmedabad','IIM Calcutta','Xavier Labour Research Institute, Jamshedpur','Faculty of Management Studies, Delhi','IIM Kozhikode','IIM Indore','SP Jain Institute of Management and Research','Management Development Institute (MDI), Gurgaon','Indian Institute of Foreign Trade (IIFT), New Delhi'	,'International Management Institute, New Delhi',	'Narsee Monjee Institute of Management Studies','National Institute of Industrial Engineering','Institute of Management Technology	Jamnalal', 'Bajaj Institute of Management Studies']
SkillDict = ['tool', 'technology', 'skill']
hobbiesDict = [item.lower() for item in hobbiesDict]
AchievementDict = [item.lower() for item in AchievementDict]
SkillDict = [item.lower() for item in SkillDict]


dict10 = ['SSC','10th','X','S.S.C','Senior Secondary']
dict12 = ['HSC', 'XII', '12th', 'H.S.C', 'Higher Secondary']
dictPg = ['PGDM','MBA','M.E.','M.B.A','M.tech','M.tech','Master']
dictBachl = ['BE', 'B.E.', 'B.tech','B.Tech', 'B.Sc.','BSc', 'Bachelor','B.E','BE.','B.Pharm','B.Sc','Engineer','Engg']
hobbiesDict = ['Hobbies', 'Hobby', 'interest']
AchievementDict = ['Achievement', 'Achievements']
skillDict = ['Statistic','Data','Min','Hypothesis','Learn','Machine','Trend','Analytic','Publish',
 'Big',
 'Advanced',
 'Algorithm',
 'Text',
 'Social',
 'Java',
 'Python',
 'R',
 'SAS',
 'Hadoop',
 'Predictive',
 'Model',
 'Communicat',
 'Domain',
 'Research',
 'Focus',
 'Ethic',
 'Matur',
 'Manage',
 'Motivat',
 'Detail',
 'Attention',
 'SPSS','php','hadoop','sql','java']

dict10 = [item.lower() for item in dict10]
dict12 = [item.lower() for item in dict12]
dictPg = [item.lower() for item in dictPg]
dictBachl = [item.lower() for item in dictBachl]
hobbiesDict = [item.lower() for item in hobbiesDict]
AchievementDict = [item.lower() for item in AchievementDict]


rhobby = ''
rAchievement = ''
rSkill = ''

from docx import Document

def getResume(pathName,docName):
    
    document = Document(join(pathName,docName))
    
    resume = []
    for p in document.paragraphs:
        t = p.text.lower()
        t = t.encode('ascii','ignore')
        resume.append(t)
    
    #    resume =  resume.append(str())
    try:
        tables = document.tables[0]
        #for table in tables:
        #    print table.cell(1,4).text
        rowCount =  len(tables.rows);
        colCount = len(tables.columns);
        
        data = []
        
        for irx, row in enumerate(tables.rows):
            temp = []
            for icx, cell in enumerate(row.cells):
                temp.append(cell.text.lower())
            data.append(temp)
        
        colNameIndx = -1
        marksIndx = -1
        degreeIndx = -1
        yearIndx = -1
        
        colList = []
        for i in xrange(0,rowCount):
            col = [row[i] for row in data];
            colList = col
            col = ' '.join(col)
            if re.findall(r'\d\d\d\d', " ".join(colList)) or re.findall(r'\d\d\d\d\s*-\s*\d\d\d\d', " ".join(colList)):
                yearIndx = i;
            elif any(x in col for x in dict10) or any(x in col for x in dict12) or any(x in col for x in dictPg) or any(x in col for x in dictBachl):
                degreeIndx = i;
                i10row = -1;
                i12row = -1;
                ipgrow = -1;
                iBerow = -1;
                for j in xrange(0,len(colList)):
                    temp = colList[j]
                    if any([x for x in dict12 if x in temp]):
                        i12row = j
                    elif any([x for x in dict10 if x in temp]):
                        i10row = j
                    elif any([x for x in dictPg if x in temp]):
                        ipgrow = j
                    elif any([x for x in dictBachl if x in temp]):
                        iBerow = j
                
            elif any(x in col for x in CollegeGradDict) or any(x in col for x in CollegePostGradDict):
                colNameIndx = i;
            elif re.findall(r'\d+.\d+%+', " ".join(col)) or re.findall(r'\d+.\d+\s*\/\s*\d+.\d+', " ".join(col)):
                marksIndx = i;
        
#        print yearIndx,degreeIndx,colNameIndx,marksIndx
#        print i10row,i12row,ipgrow,iBerow
        
        
        temp = [];
        temp.append(docName);
        
        if i10row>= 0 and colNameIndx >=0:
            temp.append(data[i10row][colNameIndx]);
            temp.append(data[i10row][marksIndx]);
        else:
            temp.append("NA")
            temp.append("NA")
            
        if i10row>= 0 and colNameIndx >=0:
            temp.append(data[i12row][colNameIndx]);
            temp.append(data[i12row][marksIndx]);
        else:
            temp.append("NA")
            temp.append("NA")
        
        if i10row>= 0 and colNameIndx >=0:
            temp.append(data[iBerow][colNameIndx]);
            temp.append(data[iBerow][marksIndx]);
        else:
            temp.append("NA")
            temp.append("NA")
        
        if i10row>= 0 and colNameIndx >=0:
            temp.append(data[ipgrow][colNameIndx]);
            temp.append(data[ipgrow][marksIndx]);
        else:
            temp.append("NA")
            temp.append("NA")
            
    except:
        b10 = False;
        b12 = False;
        bBE = False;
        bPG = False;
        r10Marks = 0;
        r12Marks = 0;
        rBEMarks = 0;
        rPGMarks = 0;
        for line in resume:
            
    #        print line
            if (any(x in line for x in dict10) or b10) and not(any(x in line for x in dict12)):
                b10 = True;
                r10Marks = re.findall(r'\d+.\d+%+', line)
                if not r10Marks:
                    continue;
                else:
                    b10 = False;
                    b12 = False;
                    bBE = False;
                    bPG = False;                
                    continue;
                
            if (any(x in line for x in dict12) or b12):
                b12 = True;
                r12Marks = re.findall(r'\d+.\d+%+', line)
                if not r12Marks:
                    continue;
                else:
                    b10 = False;
                    b12 = False;
                    bBE = False;
                    bPG = False;                
                    continue;
        
            if (any(x in line for x in dictBachl) or bBE):
                bBE = True;
                rBEMarks = re.findall(r'\d+\.\d+\/\d+\d+', line)
                if not rBEMarks:
                    continue;
                else:
                    b10 = False;
                    b12 = False;
                    bBE = False;
                    bPG = False;                
                    continue;
    
            if (any(x in line for x in dictPg) or bPG):
                bPG = True;
                rPGMarks = re.findall(r'\d+\.\d+\/\d+\.\d+', line)
                if not rPGMarks:
                    continue;
                else:
                    b10 = False;
                    b12 = False;
                    bBE = False;
                    bPG = False;                
                    continue;   
            
        temp = [];
        temp.append(docName);
        
        temp.append("NA");
        temp.append(r10Marks);
        temp.append("NA");
        temp.append(r12Marks);
        temp.append("NA");
        temp.append(rBEMarks);
        temp.append("NA");
        temp.append(rPGMarks);

    rhobby = "NA"
    rAchievement = "NA"
    rSkill = "NA"
    bHobby = False;
    bAchv = False;
    bSkill = False;
    
    for section in resume:
        
        mybool =  any(x in section for x in hobbiesDict)
        if mybool or bHobby:
            if len(section.split())<=2:
                bHobby = True;
            else:
                bHobby = False;
            rhobby = section + rhobby
                
            
        mybool =  any(x in section for x in AchievementDict)
        if mybool or bAchv:
            if len(section.split())<=2:
                bAchv = True;
            else:
                bAchv = False;
            rAchievement = section + rAchievement
        
        tempList = [x for x in skillDict if x in section]
        mybool =  any(x in section for x in skillDict)
        if mybool:
            rSkill = " ".join(tempList)
#            print rSkill
        
    temp.append(rSkill)
    temp.append(rhobby)
    temp.append(rAchievement)
    
    return temp


#head = ['name',' 10Name',' 10marks',' 12Name',' 12marks',' beName',' beMarks',' pgName',' pgMarks',' Skill',' hobby',' Achievement']


from os import listdir # for listing all files in a dir
from os.path import isfile, join # for listing all files in a dir
import os # for getting the path of currenr folder
resumeFolderPath = 'C:/Users/ad12183/Desktop/QuestFinal/resume/'
resumeFolderPath = resumeFolderPath.replace('/','\\')


finalResumes = [];
onlyfiles = [ f for f in listdir(resumeFolderPath) if isfile(join(resumeFolderPath,f)) ]
for file_name in onlyfiles:
    if file_name.lower().endswith('.doc'):
        print "converting ", convertDoc2Docx(resumeFolderPath,file_name)
    elif file_name.lower().endswith('.docx'):
        print file_name
        finalResumes.append(getResume(resumeFolderPath,file_name))


#print getResume(resumeFolderPath,"SupreetSinghKochar[3_3]13.docx")


import xlwt
book = xlwt.Workbook()
sheet1 = book.add_sheet("Sheet1")

head = ['name',' 10Name',' 10marks',' 12Name',' 12marks',' beName',' beMarks',' pgName',' pgMarks',' Skill',' hobby',' Achievement']

for i,item in enumerate(head):
    sheet1.write(0,i,item)

for i,resumeItem in enumerate(finalResumes):
    for j,item in enumerate(resumeItem):
        sheet1.write(i+1,j,item)

book.save('resumeSummary.xls')
################# new
#resumeSections = []
#temp = ''
#for line in resume:
#    if not(line == ''):
#        temp = temp + '.' + line;
#        continue;
#    else:
#        resumeSections.append(temp)
#        temp = ''
#        
#resumeSections = [name.strip() for name in resumeSections if name.strip()]
#
#bhobby = False
#bAchievement = False
#bSkill = False
#
#for section in resumeSections:
#    
#    mybool =  any(x in section for x in hobbiesDict)
#    if (mybool or bhobby):
#        rhobby = section
#        if len(rhobby.split()) <= 2:
#            bhobby = True;
#            continue;
#        else:
#            bhobby = False;
#        
#    mybool =  any(x in section for x in AchievementDict)
#    if (mybool or bAchievement):
#        rAchievement = section
#        if len(rAchievement.split()) <= 2:
#            bAchievement = True;
#            continue;
#        else:
#            bAchievement = False
#            
#    mybool =  any(x in section for x in SkillDict)
#    if (mybool or bSkill):
#        rSkill = section
#        if len(rSkill.split()) <= 2:
#            bSkill = True;
#            continue;
#        else:
#            bSkill = False
#
#
#print 'arch', rAchievement
#print 'hoby', rhobby
#print 'skill', rSkill
##a = [name.strip() for name in resume if name.strip()]
